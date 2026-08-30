#!/usr/bin/env python3
"""Gera a versao editavel (.docx) do fichamento, com o mesmo layout do PDF.

Reaproveita as citacoes e comentarios definidos em build_fichamento.py,
de modo que os dois arquivos nunca saiam do sincronismo.
"""
import importlib.util
import pathlib

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Emu

BASE = pathlib.Path(__file__).parent
OUT_DOCX = BASE / "Fichamento_Kelita_Schulz.docx"

_spec = importlib.util.spec_from_file_location("fichamento", BASE / "build_fichamento.py")
_f = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_f)

PT = 12700  # EMU por ponto


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")       # 8 oitavos de ponto = 1pt
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tcPr.append(borders)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def write_cell(cell, texto, *, size, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.15
    run = p.add_run(texto)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    set_cell_borders(cell)


def main() -> None:
    doc = Document()

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = Emu(595.28 * PT), Emu(841.89 * PT)
    sec.left_margin = Emu(56 * PT)
    sec.right_margin = Emu(56 * PT)
    sec.top_margin = Emu(56 * PT)
    sec.bottom_margin = Emu(42 * PT)
    sec.header_distance = Emu(0)

    # Faixa institucional no cabecalho, repetida em todas as paginas
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hp.add_run().add_picture(str(_f.HEADER_PNG), width=Emu(595.28 * PT))

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("CENTRO UNIVERSITÁRIO CAMPO REAL")
    r.font.size = Pt(11)

    for rotulo, valor in _f.IDENTIFICACAO:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        rb = p.add_run(rotulo + " ")
        rb.bold = True
        rb.font.size = Pt(11)
        rv = p.add_run(valor)
        rv.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(_f.TITULO)
    r.bold = True
    r.font.size = Pt(12)

    tabela = doc.add_table(rows=1, cols=2)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False

    # 12% e 88% da largura util (595,28 - 56 - 56 = 483,28 pt)
    larguras = [Emu(58 * PT), Emu(425.28 * PT)]
    cabecalhos = ["PÁGINA", "CITAÇÃO"]

    hdr = tabela.rows[0]
    set_repeat_header(hdr)
    for cel, texto, larg in zip(hdr.cells, cabecalhos, larguras):
        cel.width = larg
        write_cell(cel, texto, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Fichamento de citacoes: o comentario continua na fonte de dados,
    # mas nao e escrito no documento.
    for pagina, citacao, _comentario in _f.LINHAS:
        linha = tabela.add_row()
        c0, c1 = linha.cells
        c0.width, c1.width = larguras
        write_cell(c0, pagina, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(c1, citacao, size=10.5)

    doc.save(OUT_DOCX)
    print(f"gerado: {OUT_DOCX.name} ({len(_f.LINHAS)} linhas)")


if __name__ == "__main__":
    main()
